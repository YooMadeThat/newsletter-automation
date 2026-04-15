"""
format_agent.py — Agent 5: Word Document Renderer

Converts the Markdown newsletter output into a formatted .docx file.
No LLM API calls — uses python-docx only. Output is deterministic.

Input:  output/newsletter_YYYYMMDD.md
Output: output/newsletter_YYYYMMDD.docx
API:    None
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import harness


class FormatAgent:
    def __init__(self, output_dir: Path, run_date: date):
        self.output_dir = output_dir
        self.run_date = run_date

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def run(self, md_path: Path) -> Path:
        """Render Markdown to Word document. Returns output .docx path."""
        markdown = md_path.read_text(encoding="utf-8")

        doc = Document()
        self._configure_page(doc)

        sections = self._split_sections(markdown)
        for section_title, section_body in sections:
            self._render_section(doc, section_title, section_body)

        # Verify insert word count — log to console only, never to the document
        insert_wc = self._count_insert_words(markdown)
        if not (
            harness.NEWSLETTER_INSERT_MIN_WORDS
            <= insert_wc
            <= harness.NEWSLETTER_INSERT_MAX_WORDS
        ):
            import sys
            print(
                f"WARNING: Newsletter insert is {insert_wc} words "
                f"(expected {harness.NEWSLETTER_INSERT_MIN_WORDS}–"
                f"{harness.NEWSLETTER_INSERT_MAX_WORDS}). Review before distributing.",
                file=sys.stderr,
            )

        output_path = self.output_dir / f"newsletter_{self.run_date.strftime('%Y%m%d')}.docx"
        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def _configure_page(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width = Cm(21)    # A4
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ------------------------------------------------------------------
    # Split Markdown into (title, body) tuples by ## headings
    # ------------------------------------------------------------------

    def _split_sections(self, markdown: str) -> list[tuple[str, str]]:
        pattern = re.compile(r"^## (.+)$", re.MULTILINE)
        matches = list(pattern.finditer(markdown))
        if not matches:
            return [("Newsletter", markdown)]

        sections = []
        for i, match in enumerate(matches):
            title = match.group(1).strip()
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(markdown)
            body = markdown[start:end].strip()
            sections.append((title, body))

        return sections

    # ------------------------------------------------------------------
    # Render one section to the document
    # ------------------------------------------------------------------

    def _render_section(self, doc: Document, title: str, body: str) -> None:
        doc.add_heading(title, level=2)

        lines = body.split("\n")
        in_table = False
        table_rows: list[list[str]] = []

        i = 0
        while i < len(lines):
            line = lines[i]

            # Detect Markdown table
            if line.strip().startswith("|") and "|" in line:
                # Accumulate table rows
                if not in_table:
                    in_table = True
                    table_rows = []
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                # Skip separator rows (e.g. |---|---|)
                if not all(re.match(r"^[-: ]+$", c) for c in cells):
                    table_rows.append(cells)
                i += 1
                continue

            # End of table block — render it
            if in_table and (not line.strip().startswith("|")):
                self._render_table(doc, table_rows)
                in_table = False
                table_rows = []

            # ITEM [N]: heading → Heading 3
            if re.match(r"ITEM \[?\d+\]?:", line.strip()):
                doc.add_heading(line.strip(), level=3)
                i += 1
                continue

            # ### sub-heading
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue

            # Italicised strongest signal line (*..*)
            italic_match = re.match(r"^\*(.+)\*$", line.strip())
            if italic_match:
                para = doc.add_paragraph()
                run = para.add_run(italic_match.group(1))
                run.italic = True
                i += 1
                continue

            # Bold newsletter insert heading (**Policy + Regulation — ...**)
            bold_match = re.match(r"^\*\*(.+)\*\*$", line.strip())
            if bold_match:
                para = doc.add_paragraph()
                run = para.add_run(bold_match.group(1))
                run.bold = True
                i += 1
                continue

            # Source line — render as hyperlink paragraph
            source_match = re.match(r"^Source:\s*(.+)$", line.strip())
            if source_match:
                url = source_match.group(1).strip()
                para = doc.add_paragraph("Source: ")
                self._add_hyperlink(para, url, url)
                i += 1
                continue

            # [FACT] and [INTERPRETATION] labels — bold the label inline
            if "[FACT]" in line or "[INTERPRETATION]" in line:
                self._render_labelled_line(doc, line.strip())
                i += 1
                continue

            # Empty line — paragraph break
            if not line.strip():
                doc.add_paragraph()
                i += 1
                continue

            # Default: normal paragraph
            doc.add_paragraph(line.strip())
            i += 1

        # Flush any trailing table
        if in_table and table_rows:
            self._render_table(doc, table_rows)

    # ------------------------------------------------------------------
    # Render a Markdown table as a Word table
    # ------------------------------------------------------------------

    def _render_table(self, doc: Document, rows: list[list[str]]) -> None:
        if not rows:
            return

        col_count = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx >= col_count:
                    break
                cell = table.cell(r_idx, c_idx)
                is_header = r_idx == 0
                self._fill_table_cell(cell, cell_text.strip(), bold=is_header)

        doc.add_paragraph()  # Space after table

    def _fill_table_cell(self, cell, text: str, bold: bool = False) -> None:
        """
        Write text into a table cell. Handles markdown link syntax [label](url)
        anywhere in the cell text — renders the label as a Word hyperlink and
        any surrounding plain text as normal runs.
        """
        para = cell.paragraphs[0]
        # Clear any default content python-docx places in the cell
        for child in list(para._p):
            para._p.remove(child)

        # Split on every [label](url) occurrence, preserving order
        parts = re.split(r"(\[[^\]]+\]\([^\)]+\))", text)
        for part in parts:
            link_match = re.fullmatch(r"\[([^\]]+)\]\(([^\)]+)\)", part)
            if link_match:
                self._add_hyperlink(para, link_match.group(2), link_match.group(1))
            elif part:
                run = para.add_run(part)
                if bold:
                    run.bold = True

    # ------------------------------------------------------------------
    # Render a line with [FACT] or [INTERPRETATION] labels bolded
    # ------------------------------------------------------------------

    def _render_labelled_line(self, doc: Document, line: str) -> None:
        para = doc.add_paragraph()
        parts = re.split(r"(\[FACT\]|\[INTERPRETATION\])", line)
        for part in parts:
            if part in ("[FACT]", "[INTERPRETATION]"):
                run = para.add_run(part + " ")
                run.bold = True
            elif part.strip():
                para.add_run(part)

    # ------------------------------------------------------------------
    # Add a hyperlink to a paragraph
    # ------------------------------------------------------------------

    def _add_hyperlink(self, para, url: str, display_text: str) -> None:
        part = para.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)
        hyperlink.append(new_run)
        para._p.append(hyperlink)

    # ------------------------------------------------------------------
    # Count words in the newsletter insert (Section 4)
    # Uses the shared harness utility for consistent counting.
    # ------------------------------------------------------------------

    def _count_insert_words(self, markdown: str) -> int:
        body = harness.extract_insert_body(markdown)
        return harness.word_count(body)
